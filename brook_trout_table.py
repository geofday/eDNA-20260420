import sqlite3
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

conn = sqlite3.connect(r'c:\repos\eDNA-20260420\edna.db')

rows = conn.execute("""
    SELECT
        COALESCE(s.site_name, s.canonical_site, 'Unknown') as location,
        s.capture_time,
        s.ph_field,
        SUM(CASE WHEN e.blast_species LIKE '%fontinalis%' OR e.species LIKE '%fontinalis%'
                 THEN r.read_count ELSE 0 END) as bt_reads,
        SUM(CASE WHEN e.blast_genus NOT IN ('Struthio','Homo')
                 THEN r.read_count ELSE 0 END) as fish_reads,
        s.sample_code,
        s.batch_id
    FROM samples s
    JOIN reads r ON r.sample_code = s.sample_code
    JOIN esvs e ON e.esv_id = r.esv_id
    WHERE e.assay = 'MiFishU'
    AND s.sample_code NOT IN ('WRMXR4RT.1')
    GROUP BY s.sample_code
    HAVING fish_reads > 0
    ORDER BY (CAST(bt_reads AS REAL) / fish_reads) DESC, s.sample_code
""").fetchall()

# Load CCRO pH lookup: {(river, month): ph_avg}, plus annual averages
ccro_raw = conn.execute("SELECT river, month, ph_avg FROM ccro_ph").fetchall()
ccro_monthly = {(river, month): ph for river, month, ph in ccro_raw}
ccro_annual = {}
for river, month, ph in ccro_raw:
    ccro_annual.setdefault(river, []).append(ph)
ccro_annual = {r: sum(v)/len(v) for r, v in ccro_annual.items()}

conn.close()

CCRO_RIVER_MAP = {
    'quashnet':    'Quashnet',
    'red brook':   'Red Brook',
    'santuit':     'Santuit',
    'mashpee':     'Mashpee',
    'coonamesset': 'Coonamesset',
    'herring':     'Herring River',
}

DATE_FMTS = ('%m/%d/%Y %H:%M', '%m/%d/%Y %I:%M:%S %p', '%m/%d/%Y', '%Y-%m-%d')

def parse_date(s):
    if not s:
        return None
    for fmt in DATE_FMTS:
        try:
            return datetime.strptime(s.strip(), fmt)
        except ValueError:
            continue
    return None

def ccro_ph(site, capture_time):
    """Return (ph_value, is_ccro) using monthly match if date available, else annual avg."""
    low = str(site).lower() if site else ''
    river = next((rv for kw, rv in CCRO_RIVER_MAP.items() if kw in low), None)
    if not river:
        return None, False
    dt = parse_date(capture_time)
    if dt:
        ph = ccro_monthly.get((river, dt.month), ccro_annual.get(river))
    else:
        ph = ccro_annual.get(river)
    return ph, True

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

title = doc.add_heading('Brook Trout Detections — All Samples', level=1)
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph(f'Generated {datetime.now().strftime("%B %d, %Y")}  |  Ranked by % Brook Trout of fish reads  |  Source: edna.db')
sub.runs[0].font.size = Pt(9)
sub.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
doc.add_paragraph()

MBTS_KEYWORDS = ['sawmill', 'school st', 'fire station', 'mbts', 'proctor', 'second pond',
                 'atwater', 'cat brook', 'golf course', 'elm street', 'elm st', 'lincoln pool']

ORANGE_KEYWORDS = ['orange river', 'orange bucket', 'black fly', 'bkack fly']

ELLSWORTH_KEYWORDS = ['ellsworth', 'union river']

DEBLOIS_KEYWORDS = ['richardson', 'beaver dam', 'northern inlet', 'northern outlet', 'northern stream']

# Sample codes with no useful site name — all Downeast Washington County
DOWNEAST_CODES = {'I866I93J.1', 'S100177.1', 'SUJI1BGP.1'}

DUCKTRAP = '4JKGTSJ5.1'

TYPO_FIXES = {
    'Bkack Fly Saeason': 'Black Fly Season',
    'CohassetnNarrows nridge': 'Cohasset Narrows Bridge',
}

def normalize_location(loc, sample_code=''):
    """Returns (site, town) tuple."""
    if not loc or loc == 'Unknown':
        if sample_code == DUCKTRAP:
            return 'Ducktrap River', 'Lincoln, ME'
        return (loc or '—'), ''
    for bad, good in TYPO_FIXES.items():
        loc = loc.replace(bad, good)
    low = loc.lower()
    if any(k in low for k in ORANGE_KEYWORDS):
        if 'whiting' not in low:
            return loc.strip(), 'Whiting, ME'
    if any(k in low for k in ELLSWORTH_KEYWORDS):
        if 'ellsworth, me' not in low:
            return loc.strip(), 'Ellsworth, ME'
    is_mbts = any(k in low for k in MBTS_KEYWORDS)
    already_mbts = 'mbts, ma' in low or 'manchester' in low
    if is_mbts and not already_mbts:
        return loc.strip(), 'MBTS, MA'
    if 'tack factory' in low or 'thb' in low:
        return loc.strip(), 'Hanover/Norwell, MA'
    if 'fb pond' in low or 'fresh brook' in low:
        return loc.strip(), 'Wellfleet, MA'
    if 'mashpee' in low:
        return loc.strip(), 'Mashpee, MA'
    if 'orleans' in low:
        return loc.strip(), 'Orleans, MA'
    if 'brewster' in low:
        return loc.strip(), 'Brewster, MA'
    if 'east machias' in low:
        return loc.strip(), 'East Machias, ME'
    if 'machias river' in low or 'machias' in low:
        return loc.strip(), 'Machias, ME'
    if any(k in low for k in DEBLOIS_KEYWORDS):
        return loc.strip(), 'Deblois, ME'
    if 'gardner lake' in low or 'meatball' in low:
        return loc.strip(), 'Downeast, ME'
    if sample_code in DOWNEAST_CODES:
        return loc.strip(), 'Downeast, ME'
    return loc.strip(), ''

headers = ['Site', 'Town / State', 'Date', 'pH', '% Brook Trout\n(of fish reads)', 'Brook Trout\nreads', 'Batch', 'Sample ID']
col_widths = [Inches(1.8), Inches(1.1), Inches(0.8), Inches(0.45), Inches(1.0), Inches(0.75), Inches(0.85), Inches(0.95)]

table = doc.add_table(rows=1, cols=8)
table.style = 'Table Grid'

hdr = table.rows[0]
for i, (h, w) in enumerate(zip(headers, col_widths)):
    cell = hdr.cells[i]
    cell.width = w
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F497D')
    tcPr.append(shd)

ccro_used = 0
for i, (loc, date, ph_field, bt, fish, sample_code, batch_id) in enumerate(rows):
    pct = bt / fish * 100 if fish else 0

    dt = parse_date(date)
    date_str = dt.strftime('%m/%d/%Y') if dt else (str(date)[:10] if date else '')

    ph = ph_field
    ph_is_ccro = False
    if ph is None:
        ph, ph_is_ccro = ccro_ph(loc, date)
        if ph_is_ccro:
            ccro_used += 1

    if ph is not None:
        ph_str = f'{ph:.2f}*' if ph_is_ccro else f'{ph:.2f}'
    else:
        ph_str = '—'

    site_str, town_str = normalize_location(str(loc) if loc else '—', sample_code)
    site_str = site_str[:55] if site_str else '—'

    row = table.add_row()
    vals = [site_str, town_str, date_str, ph_str, f'{pct:.1f}%', f'{int(bt):,}',
            batch_id or '', sample_code]
    aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
              WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
              WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT,
              WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]

    for j, (v, align) in enumerate(zip(vals, aligns)):
        cell = row.cells[j]
        cell.width = col_widths[j]
        cell.text = v
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.runs[0]
        run.font.size = Pt(9)
        if i % 2 == 1:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'E8EEF6')
            tcPr.append(shd)
        if j in (6, 7):
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        if j == 4 and pct >= 20:
            run.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x6A, 0x1A)
        elif j == 4 and pct == 0:
            run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

footnote = doc.add_paragraph('* pH from Cape Cod Rivers Observatory (CCRO) monthly average. All other pH values measured in field.')
footnote.runs[0].font.size = Pt(8)
footnote.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
footnote.runs[0].font.italic = True

outpath = r'c:\repos\eDNA-20260420\output\BrookTrout_AllSamples_v9.docx'
doc.save(outpath)
print(f"Saved: {outpath}")
print(f"Rows: {len(rows)}, CCRO pH filled: {ccro_used}")
