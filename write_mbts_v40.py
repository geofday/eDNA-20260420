import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\repos\eDNA-20260420\output\MBTS_eDNA_Report_v39.docx'
DST = r'C:\repos\eDNA-20260420\output\MBTS_eDNA_Report_v40.docx'

doc = Document(SRC)

# ── HELPER: insert paragraph after matched paragraph ─────────────────────────
def insert_after_para(doc, match_text, new_text, style='Body Text'):
    for para in doc.paragraphs:
        if match_text in para.text:
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), style.replace(' ', ''))
            pPr.append(pStyle)
            p.append(pPr)
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = new_text
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r.append(t)
            p.append(r)
            para._p.addnext(p)
            return True
    return False

# ── 1. MECHANICAL ALGAL SPECIES BLAST CORRECTIONS ────────────────────────────
# Applied throughout — paragraphs and tables.

SUBS = [
    # Halamphora coffeaeformis → Eunotia naegelii
    # BLAST shows these ESVs are an acidophilous freshwater diatom, NOT the tidal halophile.
    # This reverses the "tidal signal" interpretation at all upstream MBTS sites.
    (r'Halamphora coffeaeformis', 'Eunotia naegelii'),
    (r'\bHalamphora\b',           'Eunotia naegelii'),

    # Geminigera cryophila → Teleaulax gracilis (same Cryptophyceae family; still estuarine)
    (r'Geminigera cryophila',     'Teleaulax gracilis'),
    (r'\bGeminigera\b',           'Teleaulax'),

    # Pseudopedinella elastica → Florenciella sp. (chrysophyte reclassification)
    (r'Pseudopedinella elastica', 'Florenciella sp.'),
    (r'\bPseudopedinella\b',      'Florenciella'),

    # Aureoumbra lagunensis → Veerella sp.
    (r'Aureoumbra lagunensis',    'Veerella sp.'),
    (r'\bAureoumbra\b',           'Veerella'),

    # Nanofrustulum shiloi → Fragilaria construens (Fire Station benthic diatom)
    (r'Nanofrustulum shiloi',     'Fragilaria construens'),
    (r'\bNanofrustulum\b',        'Fragilaria'),

    # Cyclotella cryptica → Stephanocyclus meneghinianus (generic reclassification)
    (r'Cyclotella cryptica',      'Stephanocyclus meneghinianus'),

    # Skeletonema pseudocostatum → Skeletonema menzelii
    (r'Skeletonema pseudocostatum', 'Skeletonema menzelii'),

    # Ostreococcus tauri → Ostreococcus sp. (species uncertain at 23S resolution)
    (r'Ostreococcus tauri',       'Ostreococcus sp.'),

    # Version bump
    (r'MBTS_eDNA_Report_v39',     'MBTS_eDNA_Report_v40'),
    (r'\bv39\b',                  'v40'),
]

def apply_subs(text):
    for pattern, replacement in SUBS:
        text = re.sub(pattern, replacement, text)
    return text

def patch_paragraph(para, fn):
    for run in para.runs:
        new = fn(run.text)
        if new != run.text:
            run.text = new
    for t in para._p.iter(qn('w:t')):
        old = t.text or ''
        new = fn(old)
        if new != old:
            t.text = new

def patch_all(doc, fn):
    for para in doc.paragraphs:
        patch_paragraph(para, fn)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    patch_paragraph(para, fn)

patch_all(doc, apply_subs)

# ── 2. ELM STREET: RETRACT HALAMPHORA TIDAL INTERPRETATION ───────────────────
# v38 cited Halamphora coffeaeformis at Elm Street (~5%) as a tidal/brackish signal.
# BLAST identifies this ESV as Eunotia naegelii — a strictly freshwater acid diatom.
# Insert a correction paragraph. Anchor on any sentence about Eunotia (was Halamphora)
# in the Elm Street section — fall back to anchoring on the Elm Street header text.

ELM_CORRECTION = (
    'BLAST correction — algal community: The ESV previously identified as Halamphora '
    'coffeaeformis at Elm Street (~5% of 23S reads) is re-identified by BLAST as '
    'Eunotia naegelii — an acidophilous freshwater diatom characteristic of soft, '
    'low-pH waters. Eunotia naegelii is not a brackish or tidal indicator; it is common '
    'in nutrient-poor streams with dissolved organic inputs. Any prior interpretation of '
    'a tidal influence at Elm Street based on Halamphora is withdrawn. The 23S algal '
    'community at Elm Street is consistent with a clean, soft-water freshwater stream reach.'
)

# Para [190] in v38: "Halamphora coffeaeformis at 5.0% reflects the tidal character..."
# After SUBS in v40, "Halamphora coffeaeformis" → "Eunotia naegelii", so anchor on the rest.
anchored = insert_after_para(doc, 'reflects the tidal character of this lower-reach Sawmill Brook site', ELM_CORRECTION, 'Body Text')
if not anchored:
    anchored = insert_after_para(doc, 'Sawmill Brook Elm Street in March 2025', ELM_CORRECTION, 'Body Text')
if not anchored:
    print('WARNING: Elm Street anchor not found — Halamphora correction not inserted')

# ── 3. SECOND POND: STEPHANODISCUS EUTROPHICATION + VEERELLA NOTE ─────────────
# Stephanodiscus minutulus dominates (22.5%) — classical eutrophic phosphorus indicator.
# Veerella sp. (was Aureoumbra lagunensis, 18.2%) — flagellate normally estuarine/coastal.
# Insert after the Carp paragraph added in v39.
STEPHANODISCUS_TEXT = (
    'Algal BLAST update — Second Pond: Stephanodiscus minutulus dominates the 23S signal '
    '(22.5% of reads, ~4,500 reads). Stephanodiscus is a classical eutrophic indicator '
    'diatom — it thrives where dissolved phosphorus is elevated and the water column is '
    'turbid. Stephanodiscus niagarae is also present. This eutrophication signature '
    'directly corroborates the Common Carp detection (30% of fish reads): Carp '
    'bottom-rooting releases phosphorus from sediment and drives the turbidity and '
    'nutrient conditions Stephanodiscus requires. Veerella sp. (formerly Aureoumbra '
    'lagunensis, 18.2%) is a flagellate most often recorded in estuarine and coastal '
    'settings; its repeated detection in a pond context warrants continued monitoring. '
    'The combined signal — eutrophic diatoms, invasive Carp, and Fathead Minnow — '
    'presents a consistent and serious habitat degradation picture at Second Pond.'
)

anchored = insert_after_para(
    doc,
    'flagged to MassWildlife and the town conservation commission',
    STEPHANODISCUS_TEXT,
    'Body Text'
)
if not anchored:
    print('WARNING: Second Pond (Stephanodiscus) anchor not found')

# ── 4. FIRE STATION: CONFIRMED TIDAL DIATOM ASSEMBLAGE ───────────────────────
# BLAST corrects Nanofrustulum shiloi → Fragilaria construens (~9% combined);
# Entomoneis sp. → Entomoneis umbratica (~8%); Haslea avium, Nitzschia supralitorea confirmed.
FIRE_STATION_ALGAE = (
    'Algal BLAST update — Fire Station: The 23S diatom assemblage confirms the Fire Station '
    'as the most estuarine-influenced MBTS site. Navicula sp. dominates (43.7%). '
    'Fragilaria construens (BLAST correction from Nanofrustulum shiloi, ~9% combined) '
    'and Entomoneis umbratica (BLAST correction from Entomoneis sp., ~8% combined) are '
    'benthic diatoms common in tidal and brackish habitats. Haslea avium (2.2%) and '
    'Nitzschia supralitorea (0.9%) are coastal diatoms whose presence at this site '
    'confirms tidal salinity influence. The Fire Station algal suite is fully consistent '
    'with the tidal limit of Sawmill Brook at Manchester Harbor.'
)

anchored = insert_after_para(
    doc,
    'late-summer feeding behavior in Manchester Harbor',
    FIRE_STATION_ALGAE,
    'Body Text'
)
if not anchored:
    print('WARNING: Fire Station algae anchor not found')

# ── 5. SCHOOL ST (TIDAL SAMPLE): GEMINIGERA → TELEAULAX; DINOPHYSIS FLAG ─────
# UTEVR3WT.1: Teleaulax gracilis replaces Geminigera cryophila (44.5%);
# Ostreococcus sp. ~30%; Dinophysis sp. trace — potential shellfish toxin genus.
SCHOOL_ST_TELEAULAX = (
    'Algal BLAST update — School St (UTEVR3WT.1): The dominant 23S signal previously '
    'attributed to Geminigera cryophila (44.5%) is re-identified by BLAST as Teleaulax '
    'gracilis — an estuarine cryptophyte in the same Cryptophyceae family. The ecological '
    'interpretation is unchanged: Teleaulax gracilis is a marine/coastal species confirming '
    'tidal influence at this sample point. Ostreococcus sp. (formerly O. tauri, ~30% of '
    'reads combined) is a picoeukaryote also characteristic of estuarine and coastal waters. '
    'Synechococcus spp. (~2%) and Skeletonema menzelii (marine diatom, 0.5%) reinforce the '
    'tidal signature. A trace of Dinophysis sp. (0.2%, 34 reads) is present — Dinophysis '
    'can produce lipophilic shellfish toxins (okadaic acid, dinophysistoxins). This detection '
    'is at very low read depth and does not constitute a bloom detection, but should be noted '
    'if recreational shellfish harvest occurs downstream in Manchester Harbor.'
)

anchored = insert_after_para(
    doc,
    'reinforces the tidal character of this reach',
    SCHOOL_ST_TELEAULAX,
    'Body Text'
)
if not anchored:
    anchored = insert_after_para(
        doc,
        'second Fundulus species to the watershed species list',
        SCHOOL_ST_TELEAULAX,
        'Body Text'
    )
if not anchored:
    print('WARNING: School St tidal algae anchor not found')

# ── 6. CAT BROOK: CHRYSOCHROMULINA/FLORENCIELLA + BRASENIA DETECTION ──────────
CAT_BROOK_ALGAE = (
    'Algal BLAST update — Cat Brook Loading Place (NVN8LUTP.1): Chrysochromulina sp. '
    'dominates the 23S signal (8.8%, 3,045 reads) — a mixotrophic golden alga common '
    'in soft-water lakes and streams with moderate dissolved organic carbon. Florenciella sp. '
    '(BLAST correction from Pseudopedinella elastica, ~13% combined) is a marine chrysophyte; '
    'its presence in a freshwater Cat Brook sample may reflect tidal connectivity at the '
    'loading place or barcode similarity to a freshwater relative. Brasenia schreberi '
    '(water shield, 0.6%, 225 reads) chloroplast DNA is detected — a native aquatic '
    'macrophyte consistent with the forested, boggy character of the upper Cat Brook '
    'catchment. Rhopalodia inflata (nitrogen-fixing diatom, 0.2%) is present, as at other '
    'MBTS 23S sites. Eunotia naegelii (BLAST correction from Halamphora coffeaeformis, '
    'trace) confirms the freshwater, soft-water character of this reach.'
)

# Para [137]: "Cat Brook / Forest Landing in October is a warm-water pond community..."
anchored = insert_after_para(
    doc,
    'Cat Brook / Forest Landing in October',
    CAT_BROOK_ALGAE,
    'Body Text'
)
if not anchored:
    anchored = insert_after_para(doc, 'NVN8LUTP.1', CAT_BROOK_ALGAE, 'Body Text')
if not anchored:
    print('WARNING: Cat Brook algae anchor not found')

# ── 7. GOLF COURSE: CHARACEAE / NITELLOPSIS UPDATE ───────────────────────────
GOLF_COURSE_STONEWORT = (
    'Stonewort BLAST update — Lower Golf Course / Sawmill Brook: Native Characeae dominates '
    'the 23S signal (32.8% of reads) — the highest aquatic macrophyte detection across all '
    'MBTS sites. This is a stream reach (lower golf course boundary, Sawmill Brook), so the '
    'Characeae signal reflects benthic stonewort beds or adjacent wetland margin vegetation '
    'in contact with flowing water. A trace of Nitellopsis obtusa (1.0%) is also detected. '
    'Nitellopsis obtusa is an invasive stonewort from Eurasia established in multiple '
    'Massachusetts waterbodies; at 1.0% read abundance a self-sustaining population cannot '
    'be confirmed from eDNA alone, but a physical survey of the stream margin is warranted. '
    'The dominant native Characeae signal is a positive habitat indicator consistent with '
    'oligotrophic, clear-water conditions and Brook Trout connectivity at this reach.'
)

anchored = insert_after_para(
    doc,
    'Mallomonas at 3.0% provides a cold-water oligotrophic baseline',
    GOLF_COURSE_STONEWORT,
    'Body Text'
)
if not anchored:
    anchored = insert_after_para(
        doc,
        'lower golf course',
        GOLF_COURSE_STONEWORT,
        'Body Text'
    )
if not anchored:
    print('WARNING: Golf Course stonewort anchor not found')

# ── 8. SCHOOL ST FRESHWATER SAMPLE: ACANTHOCERAS + DINOBRYON NOTE ────────────
# UG79PNJS.1: Acanthoceras zachariasii 35% (freshwater); Dinobryon 7.6% (oligotrophic);
# Fontinalis antipyretica ~8% (aquatic moss). Distinct from the tidal UTEVR3WT.1 sample.
SCHOOL_ST_FRESH = (
    'Algal BLAST update — Sawmill Brook School St (UG79PNJS.1): This sample, collected at '
    'a different tidal state from UTEVR3WT.1, shows a predominantly freshwater algal community. '
    'Acanthoceras zachariasii dominates (35%, 1,305 reads) — a freshwater centric diatom '
    'most common in temperate lakes and streams. Dinobryon sociale (7.6%) is a colonial '
    'chrysophyte characteristic of oligotrophic, clear waters — consistent with freshwater '
    'input from the upper Sawmill Brook reach. Fontinalis antipyretica (aquatic moss, ~8% '
    'combined across ESVs) is detected via chloroplast eDNA — the first macrophyte moss '
    'detection in the MBTS dataset. Its presence confirms submerged substrate suitable for '
    'invertebrate habitat and Bank Swallow foraging. The contrast between UG79PNJS.1 '
    '(freshwater, oligotrophic) and UTEVR3WT.1 (estuarine, Teleaulax/Ostreococcus) at the '
    'same site illustrates the tidal oscillation the School St reach experiences.'
)

anchored = insert_after_para(
    doc,
    'School Street sits in the tidal reach of Sawmill Brook',
    SCHOOL_ST_FRESH,
    'Body Text'
)
if not anchored:
    print('WARNING: School St freshwater algae anchor not found')

# ── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(DST)
print(f'Saved: {DST}')

print()
print('Algal BLAST changes applied (v39 → v40):')
print('  - Halamphora coffeaeformis → Eunotia naegelii throughout (freshwater, NOT tidal)')
print('  - Geminigera cryophila → Teleaulax gracilis (same family; estuarine interpretation unchanged)')
print('  - Pseudopedinella elastica → Florenciella sp.')
print('  - Aureoumbra lagunensis → Veerella sp.')
print('  - Nanofrustulum shiloi → Fragilaria construens')
print('  - Cyclotella cryptica → Stephanocyclus meneghinianus')
print('  - Skeletonema pseudocostatum → Skeletonema menzelii')
print('  - Ostreococcus tauri → Ostreococcus sp.')
print('  - Elm Street: tidal Halamphora interpretation retracted; Eunotia = freshwater signal')
print('  - Second Pond: Stephanodiscus eutrophication + Veerella note added')
print('  - Fire Station: Fragilaria construens + Entomoneis umbratica tidal suite confirmed')
print('  - School St (UTEVR3WT.1): Teleaulax gracilis confirmed; Dinophysis sp. trace flagged')
print('  - School St (UG79PNJS.1): Freshwater Acanthoceras/Dinobryon; Fontinalis moss detected')
print('  - Cat Brook: Chrysochromulina dominant; Brasenia schreberi macrophyte detected')
print('  - Golf Course: Characeae 32.8% native confirmed; Nitellopsis obtusa 1.0% trace flagged')
