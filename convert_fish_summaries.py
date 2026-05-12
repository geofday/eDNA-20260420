import re, os, sys
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_DEFAULT_INPUT = 'output/fish_summaries_JVB3988_JVB4279_JVB4678_JVB4846_JVB4981_JVB5403.md'
INPUT = sys.argv[1] if len(sys.argv) > 1 else _DEFAULT_INPUT
OUTPUT_BASE = os.path.splitext(INPUT)[0]

def versioned(base):
    i = 1
    while os.path.exists(f'{base}_v{i}.docx'):
        i += 1
    return f'{base}_v{i}.docx'

def set_font(run, bold=False, italic=False, size=11):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def parse_inline(para, text, bold_override=False, size=11):
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|H<sub>[^<]+</sub>|_[^_]+_)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            set_font(r, bold=True, size=size)
        elif part.startswith('*') and part.endswith('*'):
            r = para.add_run(part[1:-1])
            set_font(r, italic=True, size=size)
        elif part.startswith('_') and part.endswith('_'):
            r = para.add_run(part[1:-1])
            set_font(r, italic=True, size=size)
        elif re.match(r'H<sub>([^<]+)</sub>', part):
            m = re.match(r'H<sub>([^<]+)</sub>', part)
            sub_text = m.group(1)
            r = para.add_run('H')
            set_font(r, bold=bold_override, size=size)
            r2 = para.add_run(sub_text)
            set_font(r2, bold=bold_override, size=size)
            r2.font.subscript = True
        elif part:
            r = para.add_run(part)
            set_font(r, bold=bold_override, size=size)

def flush_table(doc, table_lines):
    rows = [l for l in table_lines if not re.match(r'\s*\|[-| :]+\|\s*$', l)]
    if not rows:
        return
    cols = max(len(r.strip().strip('|').split('|')) for r in rows)
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        for ci in range(cols):
            cell_text = cells[ci] if ci < len(cells) else ''
            tc = t.rows[ri].cells[ci]
            tc.text = ''
            p = tc.paragraphs[0]
            parse_inline(p, cell_text, bold_override=(ri == 0), size=10)

def convert():
    with open(INPUT, encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    table_lines = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Accumulate table lines
        if stripped.startswith('|'):
            table_lines.append(line)
            i += 1
            continue
        else:
            if table_lines:
                flush_table(doc, table_lines)
                table_lines = []

        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        if stripped == '---':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
            i += 1
            continue

        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue

        if stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue

        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, stripped[2:])
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline(p, stripped)
        i += 1

    if table_lines:
        flush_table(doc, table_lines)

    out = versioned(OUTPUT_BASE)
    doc.save(out)
    print(f'Saved: {out}')

convert()
