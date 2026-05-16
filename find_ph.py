import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
doc = Document(r'c:\repos\eDNA-20260420\output\MBTS_eDNA_Report_v39.docx')

print('=== All paragraphs containing pH / Vivosun / field measurement ===')
for i, para in enumerate(doc.paragraphs):
    t = para.text
    if any(kw in t.lower() for kw in ['vivosun', 'ph meter', 'ph =', 'ph:', 'field measurement', 'field ph']):
        print(f'  [{i}] {t}')

print()
print('=== Paragraphs 115-145 (around the field measurement section) ===')
for i, para in enumerate(doc.paragraphs):
    if 115 <= i <= 145:
        print(f'  [{i}] {para.text}')

print()
print('=== All tables — looking for pH values ===')
for ti, tbl in enumerate(doc.tables):
    for ri, row in enumerate(tbl.rows):
        cells = [c.text.strip() for c in row.cells]
        if any('ph' in c.lower() or any(ch in c for ch in ['6.','7.','8.','5.']) for c in cells):
            if any(c for c in cells):
                print(f'  Table {ti} Row {ri}: {cells}')
