import sys, io, copy, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from lxml import etree

SRC = r'C:\repos\eDNA-20260420\output\MBTS_eDNA_Report_v38.docx'
DST = r'C:\repos\eDNA-20260420\output\MBTS_eDNA_Report_v39.docx'

doc = Document(SRC)

# ── 1. MECHANICAL TEXT SUBSTITUTIONS ─────────────────────────────────────────
# Applied to every paragraph and every table cell.

SUBS = [
    # Lepomis — confirmed Pumpkinseed
    (r'Lepomis spp\. \(sunfish\)',      'Pumpkinseed (Lepomis gibbosus)'),
    (r'Lepomis spp\.',                  'Pumpkinseed (Lepomis gibbosus)'),
    (r'Lepomis sp\.',                   'Pumpkinseed (Lepomis gibbosus)'),
    # Micropterus — confirmed Largemouth Bass
    (r'Micropterus spp\. includes Largemouth Bass and Smallmouth Bass',
     'Micropterus salmoides (Largemouth Bass) is confirmed by BLAST. Micropterus dolomieu (Smallmouth Bass) is separately confirmed at Proctor Point'),
    (r'Micropterus spp\.',              'Largemouth Bass (Micropterus salmoides)'),
    (r'Micropterus sp\.',               'Largemouth Bass (Micropterus salmoides)'),
    # Ameiurus — confirmed Brown Bullhead
    (r'Ameiurus spp\. includes Brown Bullhead, Yellow Bullhead, and Black Bullhead',
     'Ameiurus nebulosus (Brown Bullhead) is confirmed by BLAST across all MBTS sites'),
    (r'Ameiurus spp\.',                 'Brown Bullhead (Ameiurus nebulosus)'),
    (r'Ameiurus sp\.',                  'Brown Bullhead (Ameiurus nebulosus)'),
    (r'Bullhead sp\. \(likely Brown Bullhead, Ameiurus nebulosus\)', 'Brown Bullhead (Ameiurus nebulosus)'),
    (r'Bullhead sp\.',                  'Brown Bullhead (Ameiurus nebulosus)'),
    # Cottidae — BLAST confirms Cottus sp.
    (r'Cottidae spp\. \(sculpin\) includes Slimy Sculpin and Mottled Sculpin',
     'Cottus sp. (Sculpin) is confirmed by BLAST; likely Slimy Sculpin (Cottus cognatus) or Mottled Sculpin (Cottus bairdi)'),
    (r'Cottidae spp\.',                 'Sculpin (Cottus sp.)'),
    (r'Cottidae sp\.',                  'Sculpin (Cottus sp.)'),
    (r'Sculpin \(Cottidae\)',           'Sculpin (Cottus sp.)'),
    # Brook Trout — now BLAST-confirmed to species at 4 sites
    (r'Brook Trout \(Salvelinus spp\.\)', 'Brook Trout (Salvelinus fontinalis)'),
    (r'Salvelinus spp\.',               'Salvelinus fontinalis'),
    # Largemouth Bass spp in tables
    (r'Micropterus — Bass',             'Largemouth Bass (Micropterus salmoides)'),
    (r'Non-native — competing with Brook Trout',
     'Largemouth Bass — non-native in this watershed context; competes with Brook Trout'),
    # Update report title version
    (r'MBTS_eDNA_Report_v38',           'MBTS_eDNA_Report_v39'),
    (r'v38',                            'v39'),
]

def apply_subs(text):
    for pattern, replacement in SUBS:
        text = re.sub(pattern, replacement, text)
    return text

def patch_run_text(run, fn):
    new = fn(run.text)
    if new != run.text:
        run.text = new

def patch_paragraph(para, fn):
    for run in para.runs:
        patch_run_text(run, fn)
    # Also patch any w:t nodes directly for runs that span XML nodes
    for t in para._p.iter(qn('w:t')):
        old = t.text or ''
        new = fn(old)
        if new != old:
            t.text = new

def patch_all(doc, fn):
    for para in doc.paragraphs:
        patch_paragraph(para, fn)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    patch_paragraph(para, fn)

patch_all(doc, apply_subs)

# ── 2. SALMO SPP → BROWN TROUT at MBTS #3 ────────────────────────────────────
# Find the MBTS #3 sentinel paragraph and update the Salmo note.

def patch_salmo_mbts3(doc):
    in_mbts3 = False
    for para in doc.paragraphs:
        txt = para.text
        if 'MBTS #3' in txt or 'MBTS # 3' in txt or 'JVB4846' in txt:
            in_mbts3 = True
        if in_mbts3 and ('Salmo spp.' in txt or 'Salmo trutta' in txt):
            for t in para._p.iter(qn('w:t')):
                t.text = re.sub(
                    r'Salmo spp\. at 1\.2% is genus-level only[^.]+\[T1\]\.',
                    'Salmo trutta (Brown Trout) confirmed by BLAST at 1.2% (267 reads) — non-native, stocked in many Massachusetts streams. No Atlantic Salmon presence at this site.',
                    t.text or ''
                )
            break

patch_salmo_mbts3(doc)

# ── 3. SECOND POND — ADD COMMON CARP ─────────────────────────────────────────
# Insert a new paragraph after the Alewife sentence in the Second Pond narrative.

def insert_after_para(doc, match_text, new_text, style='Body Text'):
    for i, para in enumerate(doc.paragraphs):
        if match_text in para.text:
            # Insert new paragraph after this one
            new_para = copy.deepcopy(para._p)
            para._p.addnext(new_para)
            # Now find the new paragraph object and update its text
            # Easier: manipulate XML directly
            from docx.oxml import OxmlElement
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), style.replace(' ', ''))
            pPr.append(pStyle)
            p.append(pPr)
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = new_text
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r.append(t)
            p.append(r)
            para._p.addnext(p)
            return True
    return False

CARP_TEXT = (
    'Common Carp (Cyprinus carpio) — BLAST-confirmed at 4,093 reads (~30% of fish reads at this site) — '
    'is the most significant new finding at Second Pond from the full BLAST run. Carp were '
    'previously unresolved (species=None) and excluded from the v38 community analysis. Their '
    'identification substantially revises the site picture: Alewife is no longer dominant at 63% '
    'but closer to 44% once Carp reads are included in the denominator. Carp are an introduced '
    'species with a well-documented capacity to degrade aquatic habitat through bottom-rooting, '
    'sediment resuspension, and uprooting of macrophytes. A 30% Carp signal in a headwater pond '
    'that also supports Alewife, Tessellated Darter, and potential Brook Trout connectivity is a '
    'serious conservation flag. Fathead Minnow (Pimephales promelas, ~5%) is a second introduced '
    'species now confirmed by BLAST — a common bait-bucket introduction. Both species should be '
    'flagged to MassWildlife and the town conservation commission.'
)

insert_after_para(
    doc,
    'Mallomonas at 3.0% provides a cold-water oligotrophic baseline',
    CARP_TEXT,
    'Body Text'
)

# ── 4. FIRE STATION — ADD BLUEFISH NOTE ──────────────────────────────────────
BLUEFISH_TEXT = (
    'Bluefish (Pomatomus saltator) detected at 167 reads (0.6%) at the Fire Station tidal limit — '
    'a marine pelagic predator whose eDNA at the tidal mouth of Sawmill Brook is consistent with '
    'late-summer feeding behavior in Manchester Harbor. Bluefish are voracious predators of '
    'Mummichog, Alewife, and other small fish common in this tidal reach. Their eDNA at the '
    'tidal limit is not unexpected in late August.'
)

insert_after_para(
    doc,
    'Mummichog 95.2% + estuarine diatoms',
    BLUEFISH_TEXT,
    'Body Text'
)

# ── 5. UPDATE T1 NOTE — LEPOMIS AND MICROPTERUS NOW RESOLVED ─────────────────
def update_t1(doc):
    for para in doc.paragraphs:
        if '[T1]' in para.text and 'Lepomis spp.' in para.text:
            for t in para._p.iter(qn('w:t')):
                txt = t.text or ''
                txt = txt.replace(
                    'Lepomis spp. (sunfish) includes Bluegill, Pumpkinseed, Redear Sunfish, and Longear Sunfish (L. megalotis), among others.',
                    'Lepomis — BLAST has now resolved all MBTS Lepomis ESVs to Lepomis gibbosus (Pumpkinseed). Pumpkinseed is native to eastern North America. No Bluegill (L. macrochirus, introduced) confirmed in this dataset.'
                )
                t.text = txt

update_t1(doc)

# ── 6. ADD SCHOOL ST STRIPED KILLIFISH NOTE ───────────────────────────────────
KILLIFISH_TEXT = (
    'BLAST update: Fundulus majalis (Striped Killifish) is now confirmed at School Street '
    '(821 reads in JVB4678 / UTEVR3WT.1). Striped Killifish is a tidal species distinct from '
    'Mummichog (F. heteroclitus), preferring sandy tidal flats and creek margins. Its confirmation '
    'at School Street reinforces the tidal character of this reach and adds a second Fundulus '
    'species to the watershed species list.'
)
insert_after_para(
    doc,
    'School Street sits in the tidal reach of Sawmill Brook',
    KILLIFISH_TEXT,
    'Body Text'
)

# ── 7. UPDATE TITLE DATE LINE ─────────────────────────────────────────────────
for para in doc.paragraphs:
    for t in para._p.iter(qn('w:t')):
        if 'April 2026' in (t.text or ''):
            t.text = t.text.replace('April 2026', 'May 2026')

# ── SAVE ──────────────────────────────────────────────────────────────────────
doc.save(DST)
print(f'Saved: {DST}')

# Summary of changes
print()
print('Changes applied:')
print('  - Lepomis spp. → Pumpkinseed (Lepomis gibbosus) throughout')
print('  - Micropterus spp. → Largemouth Bass (Micropterus salmoides) throughout')
print('  - Ameiurus spp. → Brown Bullhead (Ameiurus nebulosus) throughout')
print('  - Bullhead sp. → Brown Bullhead (Ameiurus nebulosus) throughout')
print('  - Cottidae spp. → Sculpin (Cottus sp.) throughout')
print('  - Brook Trout (Salvelinus spp.) → Brook Trout (Salvelinus fontinalis) throughout')
print('  - MBTS #3: Salmo spp. → Brown Trout (Salmo trutta) confirmed, non-native flagged')
print('  - Second Pond: Common Carp (30%) + Fathead Minnow new paragraphs added')
print('  - Fire Station: Bluefish (Pomatomus saltator) note added')
print('  - School St: Striped Killifish (Fundulus majalis) note added')
print('  - T1 note: Lepomis updated to reflect BLAST resolution')
