"""Convert ecological_narrative_JVB5776.md to a formatted Word document."""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = 'output/ecological_narrative_JVB5776.md'
OUTPUT_BASE = 'output/ecological_narrative_JVB5776'


def _versioned(base, ext='.docx'):
    v = 1
    while os.path.exists(f'{base}_v{v}{ext}'):
        v += 1
    return f'{base}_v{v}{ext}'


def add_subscript_run(para, text):
    """Add a run with subscript XML for H_f / H_a notation."""
    run = para.add_run(text)
    rpr = run._r.get_or_add_rPr()
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'subscript')
    rpr.append(vertAlign)
    run.font.size = Pt(8)
    return run


def parse_inline(para, text, bold_override=False):
    """
    Parse inline markdown: **bold**, *italic*, H<sub>f</sub>, H<sub>a</sub>.
    Writes runs to para.
    """
    # Tokenise: split on **...**, *...*, <sub>...</sub>
    pattern = re.compile(
        r'(\*\*[^*]+\*\*|\*[^*]+\*|H<sub>[^<]+</sub>)'
    )
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
            if bold_override:
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif re.match(r'H<sub>([^<]+)</sub>', part):
            m = re.match(r'H<sub>([^<]+)</sub>', part)
            sub = m.group(1)
            run = para.add_run('H')
            if bold_override:
                run.bold = True
            add_subscript_run(para, sub)
        else:
            run = para.add_run(part)
            if bold_override:
                run.bold = True


def set_heading_style(para, level):
    if level == 1:
        para.style = 'Heading 1'
    elif level == 2:
        para.style = 'Heading 2'
    elif level == 3:
        para.style = 'Heading 3'


def add_table_row(table, cells, is_header=False):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        para = cell.paragraphs[0]
        parse_inline(para, cell_text.strip(), bold_override=is_header)
        if is_header:
            for run in para.runs:
                run.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def convert():
    with open(INPUT, encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Default body font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip horizontal rules
        if re.match(r'^-{3,}$', line):
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,3})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2)
            para = doc.add_paragraph()
            set_heading_style(para, level)
            parse_inline(para, heading_text)
            i += 1
            continue

        # Table: collect all rows
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1
            # Parse columns from first row
            header_cells = [c for c in table_lines[0].split('|') if c.strip() != '']
            ncols = len(header_cells)
            table = doc.add_table(rows=0, cols=ncols)
            table.style = 'Table Grid'
            # Header row
            add_table_row(table, header_cells, is_header=True)
            # Data rows (skip separator row — all dashes)
            for tline in table_lines[2:]:
                cells = [c for c in tline.split('|') if c != '']
                if len(cells) == ncols:
                    add_table_row(table, cells)
            doc.add_paragraph()  # spacing after table
            continue

        # Bullet list
        m = re.match(r'^- (.*)', line)
        if m:
            para = doc.add_paragraph(style='List Bullet')
            parse_inline(para, m.group(1))
            i += 1
            continue

        # Bold-only line (site header like **Fish: ... | Algae: ...**)
        if line.startswith('**') and line.endswith('**'):
            para = doc.add_paragraph()
            parse_inline(para, line)
            i += 1
            continue

        # Empty line → paragraph break
        if line.strip() == '':
            i += 1
            continue

        # Regular paragraph
        para = doc.add_paragraph()
        parse_inline(para, line)
        i += 1

    out = _versioned(OUTPUT_BASE)
    doc.save(out)
    print(f'Saved: {out}')


if __name__ == '__main__':
    convert()
