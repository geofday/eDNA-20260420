import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
doc = Document(r'c:\repos\eDNA-20260420\output\MBTS_eDNA_Report_v39.docx')
print('=== ELM STREET PARAGRAPHS ===')
for i, para in enumerate(doc.paragraphs):
    if 'elm' in para.text.lower() or 'AEFBOYWK' in para.text:
        print(f'  [{i}] {repr(para.text[:130])}')
print()
print('=== CAT BROOK PARAGRAPHS ===')
for i, para in enumerate(doc.paragraphs):
    if 'cat brook' in para.text.lower() or 'NVN8LUTP' in para.text or 'cat br' in para.text.lower():
        print(f'  [{i}] {repr(para.text[:130])}')
