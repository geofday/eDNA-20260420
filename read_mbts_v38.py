import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
doc = Document(r'C:\repos\eDNA-20260420\output\MBTS_eDNA_Report_v38.docx')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{p.style.name[:18]:18s}] {p.text}")
print()
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
for t_i, tbl in enumerate(doc.tables):
    print(f"\nTable {t_i}: {len(tbl.rows)} rows x {len(tbl.columns)} cols")
    for row in tbl.rows[:3]:
        print("  ", [c.text[:40] for c in row.cells])
