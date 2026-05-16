import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DST = r'C:\repos\eDNA-20260420\output\MBTS_eDNA_Summary_v1.docx'
doc = Document()

# ── Page margins — tight to fit one page ─────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(0.6)
section.bottom_margin = Inches(0.6)
section.left_margin   = Inches(0.8)
section.right_margin  = Inches(0.8)

def heading(doc, text, level=1, size=11, bold=True, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6 if level == 1 else 3)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def body(doc, text, size=8.5, space_after=2):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    for run in p.runs:
        run.font.size = Pt(size)
    return p

def bullet(doc, text, size=8.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.15)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

# ── TITLE ─────────────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(0)
t.paragraph_format.space_after  = Pt(2)
r = t.add_run('MBTS eDNA Watershed Summary — May 2026')
r.bold = True; r.font.size = Pt(12)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(0)
sub.paragraph_format.space_after  = Pt(4)
rs = sub.add_run('Sawmill Brook / Cat Brook Watershed, Manchester-by-the-Sea MA  |  14 samples, 2024–2025  |  v40 findings')
rs.font.size = Pt(8.5); rs.italic = True

# ── 1. BROOK TROUT ────────────────────────────────────────────────────────────
heading(doc, '1.  Brook Trout (Salvelinus fontinalis)', size=10, color=(0,70,127))
body(doc,
    'Confirmed at 4 of 14 sites. Strongest signal at Below School St (4.6% of fish reads, '
    'Nov 2025 — spawning season) and Sawmill Brook at School St (2.5%, Aug 2024). Sawmill '
    'Swamp (1.0%, Aug 2024) and Below Lincoln Pool (0.6%, Jun 2025) are also positive. '
    'Brook Trout were absent from Upper Sawmill, Golf Course reach, Cat Brook, Second Pond, '
    'Atwater, MBTS #3, and Proctor Point in all samples to date. The Nov 2025 absence at '
    'Upper Sawmill during spawning season is a flag — this reach should be a functional '
    'spawning corridor. BLAST confirms all positive detections to Salvelinus fontinalis.')

# ── 2. EUTROPHICATION ─────────────────────────────────────────────────────────
heading(doc, '2.  Eutrophication', size=10, color=(0,70,127))
body(doc,
    'Second Pond is the single most degraded site in the watershed and the only eutrophication '
    'signal. Common Carp (37% of fish reads, BLAST-confirmed) and Fathead Minnow (~6%) are '
    'introduced species whose presence was unknown before BLAST. Carp bottom-rooting releases '
    'phosphorus from sediment; the 23S algal community confirms the result: Stephanodiscus '
    'minutulus (22.5%) — a classical high-phosphorus diatom — dominates, alongside Veerella sp. '
    '(18%), a eutrophic flagellate. Alewife at 42% (not 63% as v38 reported before Carp were '
    'included in the denominator) remains present but is no longer ecologically dominant. '
    'The main-stem Sawmill Brook community (Eunotia naegelii, Mallomonas, Dinobryon, '
    'Chrysochromulina) is oligotrophic throughout — the nutrient problem is geographically '
    'contained to Second Pond but ecologically critical given its headwater position feeding Cat Brook.')

# ── 3. pH ─────────────────────────────────────────────────────────────────────
heading(doc, '3.  pH', size=10, color=(0,70,127))
body(doc,
    'No field pH measurements are available in the database for any of the 14 MBTS samples. '
    'This is a significant data gap. Rhopalodia inflata (N-fixing diatom) at multiple main-stem '
    'sites and Eunotia naegelii dominance throughout indicate nitrogen-limited, soft, low-pH '
    'conditions typical of New England Brook Trout habitat — but no direct measurement confirms '
    'this. Priority sites for pH monitoring: Second Pond (eutrophication severity assessment), '
    'MBTS #3 (Brown Trout stocking context), Golf Course reach (Brook Trout spawning corridor), '
    'and Sawmill Swamp. A single seasonal Sonde deployment or grab-sample run would close this gap.')

# ── 4. SENTINEL SPECIES ───────────────────────────────────────────────────────
heading(doc, '4.  Significant and Sentinel Species', size=10, color=(0,70,127))

bullet(doc,
    'Largemouth Bass (Micropterus salmoides) — dominant across most sites: 98% at Sawmill Swamp, '
    '66% at Cat Brook, 46% at Below School St, present at Upper Sawmill, Lincoln Pool, School St. '
    'The single greatest competitive threat to Brook Trout in the watershed.')
bullet(doc,
    'Pumpkinseed (Lepomis gibbosus) — ubiquitous: 99% at Golf Course, 93% Upper Sawmill, 91% '
    'MBTS #3, 83% Lincoln Pool, 88% School St Aug. Displaces invertebrate prey throughout.')
bullet(doc,
    'Brown Trout (Salmo trutta) — MBTS #3 only (8.5%, 267 reads, BLAST-confirmed). Non-native, '
    'stocked. Competes directly with Brook Trout for territory and food.')
bullet(doc,
    'Smallmouth Bass (Micropterus dolomieu) — Proctor Point (46%, BLAST-confirmed first detection '
    'in watershed). Expanding range; direct Brook Trout competitor in cold reaches.')
bullet(doc,
    'Common Carp (Cyprinus carpio) — Second Pond only (37%). Driving eutrophication. '
    'Flag to MassWildlife.')
bullet(doc,
    'Fathead Minnow (Pimephales promelas) — Second Pond (~6%). Bait-bucket introduction; '
    'second invasive species at this site.')
bullet(doc,
    'Bluefish (Pomatomus saltator) — Fire Station tidal limit (167 reads, Aug). Marine predator; '
    'seasonal presence consistent with Manchester Harbor late-summer feeding.')
bullet(doc,
    'Striped Killifish (Fundulus majalis) — School St tidal reach (821 reads). Second Fundulus '
    'species confirmed; reinforces tidal character of this reach.')

# ── 5. HABITAT INDICATORS ─────────────────────────────────────────────────────
heading(doc, '5.  Habitat Indicators', size=10, color=(0,70,127))
body(doc,
    'Characeae stonewort (32.8% of 23S reads) at the Golf Course reach and Callitriche stagnalis '
    '(water starwort, visually confirmed + eDNA at Below School St, Upper Sawmill, Golf Course) '
    'indicate submerged aquatic vegetation — positive structure for Brook Trout and invertebrates. '
    'Fontinalis antipyretica (aquatic moss, eDNA) at School St confirms suitable submerged substrate. '
    'Nitellopsis obtusa (invasive stonewort, 1% trace) at Golf Course reach warrants a physical '
    'survey. The Elm Street algal community (Eunotia naegelii dominant) confirms a clean, '
    'soft-water freshwater reach — no tidal influence extends above School St.')

# ── Footer note ───────────────────────────────────────────────────────────────
doc.add_paragraph()
fn = doc.add_paragraph()
fn.paragraph_format.space_before = Pt(2)
fnr = fn.add_run(
    'Data source: edna.db (SQLite); MiFishU 12S (fish) + 23S (algal/macrophyte) amplicon sequencing; '
    'BLAST v40 refinements applied. All species % = reads / total fish reads (MiFishU) or total 23S '
    'reads at site. pH: no field measurements available for MBTS sites. Generated May 2026.')
fnr.font.size = Pt(7.5)
fnr.italic = True
fn.paragraph_format.space_after = Pt(0)

doc.save(DST)
print(f'Saved: {DST}')
